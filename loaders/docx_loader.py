from .base_loader import BaseLoader

try:
    from docx import Document
except ImportError:
    raise ImportError("请安装python-docx: pip install python-docx")

import os

class DocxLoader(BaseLoader):
    """Word文档加载器"""
    def load(self):
        """加载Word文档"""
        try:
            doc = Document(self.file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"

            # 提取表格数据
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            self.document = {
                "content": text,
                "metadata": {
                    **self.get_metadata(),
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(tables)
                },
                "tables": tables,
                "type": "docx"
            }
            return self.document
        except Exception as e:
            print(f"加载Word文档失败: {e}")
            return None